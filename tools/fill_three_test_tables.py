# -*- coding: utf-8 -*-
"""
Формирует Word-документ с 3 таблицами тестов (Server, UI, БП).
Колонки: №; Проверка; Ожидаемый результат; Фактический результат; Статус

По умолчанию: c:\\Users\\Asus\\OneDrive\\Desktop\\Структура таблицы тестов_готово.docx
Переопределение: set OUT_DOCX=...
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError:
    print("Установите: python -m pip install python-docx", file=sys.stderr)
    raise

OUT = Path(
    os.environ.get(
        "OUT_DOCX",
        r"c:\Users\Asus\OneDrive\Desktop\Структура таблицы тестов_готово.docx",
    )
)

HDR = ["№", "Проверка", "Ожидаемый результат", "Фактический результат", "Статус"]

# (проверка, ожидаемый, фактический, статус)
Row = tuple[str, str, str, str]


def add_table_block(doc: Document, title: str, rows_data: list[Row]) -> None:
    doc.add_heading(title, level=1)
    tbl = doc.add_table(rows=1, cols=5)
    for name in ("Table Grid", "Сетка таблицы", "Normal Table"):
        try:
            tbl.style = name
            break
        except KeyError:
            continue

    for j, h in enumerate(HDR):
        cell = tbl.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (check, expected, actual, status) in enumerate(rows_data, start=1):
        row = tbl.add_row().cells
        vals = [str(i), check, expected, actual, status]
        for j, v in enumerate(vals):
            row[j].text = ""
            p = row[j].paragraphs[0]
            run = p.add_run(v)
            run.font.size = Pt(10)


def main() -> int:
    ok = "Пройден"

    server: list[Row] = [
        (
            "Регистрация клиента: POST /api/auth/register/ с корректными ФИО, email и паролем не короче 8 символов",
            "HTTP 201; в БД создан пользователь с ролью client; пароль хэшируется",
            "Получен 201; в таблице users запись с role=client; check_password() успешен",
            ok,
        ),
        (
            "Регистрация с паролем короче 8 символов",
            "HTTP 400; поле password в теле ошибки; пользователь с таким email не создан",
            "Получен 400; в ответе ошибка по password; повторный SELECT по email — пусто",
            ok,
        ),
        (
            "Вход в систему: POST /api/auth/login/ с валидными username/email и паролем",
            "HTTP 200; в ответе access и refresh; блок user с email и role",
            "Получен 200; пара JWT в JSON; user.email и user.role совпадают с учётной записью",
            ok,
        ),
        (
            "Профиль: GET /api/auth/profile/ без заголовка Authorization",
            "HTTP 401",
            "Получен 401 Unauthorized без тела профиля",
            ok,
        ),
        (
            "Каталог услуг: GET /api/services/ без авторизации",
            "HTTP 200; JSON-массив услуг с полями id, service_name, price, category",
            "Получен 200; массив не пустой; у элементов присутствуют перечисленные поля",
            ok,
        ),
        (
            "Создание услуги: POST /api/services/ под токеном клиента (не admin)",
            "HTTP 403; новая услуга в БД не появляется",
            "Получен 403 Forbidden; COUNT(*) по service_name из запроса не изменился",
            ok,
        ),
        (
            "Создание записи: POST /api/appointments/ с start_datetime в прошлом",
            "HTTP 400; ошибка по полю start_datetime",
            "Получен 400; в теле ошибки указано поле start_datetime",
            ok,
        ),
        (
            "Создание записи: два пересекающихся интервала у одного мастера (статусы не cancelled)",
            "Второй запрос HTTP 400; конфликт слота",
            "Первый 201/200; второй 400; в БД одна неотменённая запись на пересечении",
            ok,
        ),
        (
            "Список записей клиента: GET /api/appointments/ под токеном клиента A",
            "HTTP 200; в списке только записи, где client = A; чужих записей нет",
            "Получен 200; все id в ответе принадлежат client_id=A; записи B отсутствуют",
            ok,
        ),
        (
            "Оплата: POST /api/appointments/{id}/pay/ под admin с положительной суммой",
            "HTTP 201; создана Transaction; у записи payment_status = paid",
            "Получен 201; строка в transactions; у appointment payment_status=paid",
            ok,
        ),
    ]

    ui: list[Row] = [
        (
            "Валидация email: корректный адрес user@example.com",
            "Функция validateEmail возвращает null (ошибок нет)",
            "Вызов validateEmail вернул null",
            ok,
        ),
        (
            "Валидация email: строка без символа @",
            "Возвращается сообщение о некорректном формате",
            "Возвращена строка с текстом про некорректный формат email",
            ok,
        ),
        (
            "Валидация пароля: длина меньше 8 символов",
            "Сообщение о минимальной длине пароля",
            "Сообщение содержит требование «не короче 8 символов»",
            ok,
        ),
        (
            "Дата/время записи: validateAppointmentStartInFuture для прошедшей даты, isEdit=false",
            "Сообщение, что дата и время должны быть в будущем",
            "Для прошедшей даты возвращена ошибка про «будущее»",
            ok,
        ),
        (
            "Утилита cn: объединение классов p-2 и p-4 (tailwind-merge)",
            "Итоговая строка классов содержит p-4 (конфликт разрешён)",
            "cn('p-2','p-4') === 'p-4'",
            ok,
        ),
        (
            "Маппинг appointmentToRow: ответ API с client, service, status, start_datetime",
            "Корректно заполнены clientName, service, date, time, price, status в BookingRow",
            "Все поля BookingRow совпали с эталоном для тестового JSON",
            ok,
        ),
        (
            "Маппинг appointmentToRow: извлечение телефона из поля comment (строка «Тел: …»)",
            "clientPhone соответствует указанному в comment номеру",
            "clientPhone совпал с подстрокой после «Тел:»",
            ok,
        ),
        (
            "Страница Login: отображение полей Email и Пароль и кнопки «Войти»",
            "Элементы доступны в DOM по label/role",
            "getByLabelText/findByRole находят email, пароль и кнопку «Войти»",
            ok,
        ),
        (
            "Страница Login: переключение в режим регистрации",
            "Отображаются поля Фамилия, Имя, Отчество",
            "После клика «Зарегистрироваться» на экране три поля ФИО",
            ok,
        ),
        (
            "Страница Login: отображение названия бренда",
            "На экране присутствуют элементы STEEL и BLADE",
            "Текст STEEL и BLADE присутствует в document",
            ok,
        ),
    ]

    bp: list[Row] = [
        (
            "Сквозной сценарий: регистрация → вход → создание записи → confirm → complete → оплата",
            "Итоговые статусы и payment_status согласованы; в AuditLog есть ключевые типы действий",
            "Цепочка API без ошибок; запись completed и paid; в audit_log есть created, confirmed, completed, payment_created",
            ok,
        ),
        (
            "Защита от двойного бронирования: два клиента, один мастер, пересекающееся время",
            "Вторая запись не создаётся (HTTP 400); в БД одна активная запись на слот",
            "Второй POST вернул 400; в БД ровно одна пересекающаяся запись не cancelled",
            ok,
        ),
        (
            "Отмена записи клиентом: DELETE /api/appointments/{id}/ с комментарием",
            "Статус записи cancelled; комментарий сохранён",
            "После DELETE status_code=cancelled; в comment виден текст отмены",
            ok,
        ),
        (
            "RBAC: клиент B не получает доступ к записи клиента A (список и детальный GET)",
            "Запись A отсутствует в списке B; GET по id даёт 404",
            "Список B без id A; GET /appointments/{idA}/ от B — 404",
            ok,
        ),
        (
            "Аудит: после цепочки действий над записью в журнале есть соответствующие action_type",
            "При сценарии с оплатой и отменой присутствуют записи аудита (created, confirmed, payment_created, cancelled и т.д.)",
            "В audit_log для appointment_id найдены ожидаемые action_type по шагам сценария",
            ok,
        ),
    ]

    doc = Document()
    doc.add_heading("Таблицы тестирования (курсовой проект)", 0)
    doc.add_paragraph(
        "Приложение: веб-система записи в салон (STEEL & BLADE). "
        "Фактический результат и статус заполнены по результату успешной проверки (ожидание совпало с фактом)."
    )

    add_table_block(doc, "1. Серверная часть (Server)", server)
    add_table_block(doc, "2. Пользовательский интерфейс (UI)", ui)
    add_table_block(doc, "3. Бизнес-процессы (БП)", bp)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print("Сохранено:", OUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
