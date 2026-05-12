# -*- coding: utf-8 -*-
"""Заполняет первую таблицу в docx шаблоне тестами (проект Kursowaya-test)."""
from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("pip install python-docx", file=sys.stderr)
    raise

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Структура таблицы тестов.docx"


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)


def main() -> int:
    if not DOCX.is_file():
        print(f"Нет файла: {DOCX}", file=sys.stderr)
        return 1

    doc = Document(str(DOCX))
    if not doc.tables:
        # Создаём таблицу, если в шаблоне её не было в теле документа
        doc.add_paragraph("Таблица тестирования приложения «STEEL & BLADE» (курсовой проект).")
        tbl = doc.add_table(rows=1, cols=8)
        for style_name in ("Table Grid", "Сетка таблицы", "Normal Table"):
            try:
                tbl.style = style_name
                break
            except KeyError:
                continue
        hdr = [
            "№",
            "Вид теста",
            "Идентификатор / название",
            "Проверяемый объект",
            "Входные данные / действия",
            "Ожидаемый результат",
            "Фактический результат",
            "Статус",
        ]
        for i, h in enumerate(hdr):
            set_cell_text(tbl.rows[0].cells[i], h)
        start_row = 1
        table = tbl
    else:
        table = doc.tables[0]
        start_row = 1  # строка 0 — заголовок, не трогаем если уже заполнен
        # если первая строка пустая — заполним заголовки
        if not any(c.text.strip() for c in table.rows[0].cells):
            hdr = [
                "№",
                "Вид теста",
                "Идентификатор / название",
                "Проверяемый объект",
                "Входные данные / действия",
                "Ожидаемый результат",
                "Фактический результат",
                "Статус",
            ]
            ncols = len(table.rows[0].cells)
            for i in range(min(ncols, len(hdr))):
                set_cell_text(table.rows[0].cells[i], hdr[i])

    # Данные тестов (Server / UI / БП) под стек Django REST + React
    rows: list[list[str]] = []

    # --- Server (10) ---
    server = [
        ("S-01", "Server", "Регистрация клиента", "POST /api/auth/register/", "Корректные ФИО, email, пароль ≥8", "HTTP 201, пользователь создан, роль client", "", ""),
        ("S-02", "Server", "Регистрация: короткий пароль", "POST /api/auth/register/", "Пароль < 8 символов", "HTTP 400, пользователь не создан", "", ""),
        ("S-03", "Server", "Вход JWT", "POST /api/auth/login/", "Валидные учётные данные", "HTTP 200, access и refresh в ответе", "", ""),
        ("S-04", "Server", "Профиль без токена", "GET /api/auth/profile/", "Без заголовка Authorization", "HTTP 401", "", ""),
        ("S-05", "Server", "Список услуг", "GET /api/services/", "Без авторизации", "HTTP 200, массив услуг", "", ""),
        ("S-06", "Server", "Создание услуги не админом", "POST /api/services/", "Токен клиента, тело новой услуги", "HTTP 403", "", ""),
        ("S-07", "Server", "Запись в прошлом", "POST /api/appointments/", "start_datetime в прошлом", "HTTP 400, ошибка по полю времени", "", ""),
        ("S-08", "Server", "Пересечение слотов мастера", "POST /api/appointments/", "Два запроса на одного мастера с перекрытием по времени", "Второй запрос HTTP 400", "", ""),
        ("S-09", "Server", "Список записей клиента", "GET /api/appointments/", "Токен клиента A", "Только записи клиента A", "", ""),
        ("S-10", "Server", "Оплата записи", "POST /api/appointments/{id}/pay/", "Токен admin, сумма > 0", "HTTP 201, Transaction, payment_status=paid", "", ""),
    ]
    # --- UI (10) ---
    ui = [
        ("U-01", "UI", "Валидация email (валид)", "formValidation.validateEmail", "user@example.com", "null (ошибок нет)", "", ""),
        ("U-02", "UI", "Валидация email (невалид)", "formValidation.validateEmail", "строка без @", "Сообщение об ошибке формата", "", ""),
        ("U-03", "UI", "Валидация пароля", "formValidation.validatePassword", "Пароль 5 символов", "Сообщение «не короче 8»", "", ""),
        ("U-04", "UI", "Дата записи в прошлом", "validateAppointmentStartInFuture", "Дата на неделю назад, isEdit=false", "Сообщение о необходимости будущего времени", "", ""),
        ("U-05", "UI", "Утилита cn / twMerge", "lib/utils.cn", "Классы p-2 и p-4", "Итог p-4", "", ""),
        ("U-06", "UI", "appointmentToRow: основные поля", "appointmentsApi", "JSON записи с клиентом, услугой, статусом", "Корректные clientName, service, price, status", "", ""),
        ("U-07", "UI", "appointmentToRow: телефон из comment", "appointmentsApi", "comment с строкой «Тел: …»", "clientPhone извлечён", "", ""),
        ("U-08", "UI", "Страница Login: поля", "Login.tsx", "Рендер компонента", "Видны Email, Пароль, кнопка «Войти»", "", ""),
        ("U-09", "UI", "Login: режим регистрации", "Login.tsx", "Переключение на регистрацию", "Поля Фамилия, Имя, Отчество", "", ""),
        ("U-10", "UI", "Login: бренд", "Login.tsx", "Заголовок окна", "Отображается STEEL / BLADE", "", ""),
    ]
    # --- БП (5) ---
    bp = [
        ("B-01", "БП", "Жизненный цикл записи", "API + БД", "Регистрация→логин→запись→confirm→complete→pay", "Статусы и оплата согласованы, записи в AuditLog", "", ""),
        ("B-02", "БП", "Двойное бронирование", "API", "Два клиента, один мастер, пересекающееся время", "Вторая запись отклонена", "", ""),
        ("B-03", "БП", "Отмена клиентом", "API", "DELETE записи с комментарием", "Статус cancelled, комментарий сохранён", "", ""),
        ("B-04", "БП", "RBAC: чужая запись", "API", "Клиент B запрашивает запись клиента A", "HTTP 404 / нет в списке", "", ""),
        ("B-05", "БП", "Журнал аудита", "AuditLog", "Цепочка действий над записью", "Есть created, confirmed, payment, cancelled по сценарию", "", ""),
    ]

    for batch in (server, ui, bp):
        rows.extend(batch)

    ncols = len(table.rows[0].cells)
    # удалить старые строки данных (кроме заголовка)
    while len(table.rows) > 1:
        tr = table.rows[1]._tr
        tr.getparent().remove(tr)

    for idx, r in enumerate(rows, start=1):
        cells = table.add_row().cells
        vals = [
            str(idx),
            r[1],
            r[2],
            r[3],
            r[4],
            r[5],
            r[6],
            r[7],
        ]
        for j in range(min(ncols, len(vals))):
            set_cell_text(cells[j], vals[j])

    out = DOCX.with_name("Структура таблицы тестов_заполнено.docx")
    doc.save(str(out))
    print("Saved:", out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
