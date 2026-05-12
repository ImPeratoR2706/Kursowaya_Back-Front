import argparse
from docx import Document


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx_path")
    parser.add_argument("--start", type=int, default=0)
    parser.add_argument("--end", type=int, default=50)
    parser.add_argument("--out", default="")
    args = parser.parse_args()

    doc = Document(args.docx_path)
    start = max(0, args.start)
    end = min(len(doc.paragraphs), args.end)
    lines: list[str] = []
    for i in range(start, end):
        text = doc.paragraphs[i].text.strip()
        if text:
            lines.append(f"{i}: {text}")

    if args.out:
        with open(args.out, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
    else:
        for line in lines:
            print(line)


if __name__ == "__main__":
    main()
